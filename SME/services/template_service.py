import os
from docx import Document
from app_config import TEMPLATE_FOLDER

class TemplateService:
    def fill_template(self, template_name, data: dict, subfolder=""):
        path = os.path.join(TEMPLATE_FOLDER, subfolder, template_name)
        doc = Document(path)

        for p in doc.paragraphs:
            for key, value in data.items():
                if f"[{key}]" in p.text:
                    p.text = p.text.replace(f"[{key}]", str(value))

        # Also replace in tables (if needed)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        if f"[{key}]" in cell.text:
                            cell.text = cell.text.replace(f"[{key}]", str(value))

        return doc
